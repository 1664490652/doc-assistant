import os
import re
import docx
import pdfplumber

class DocumentParser:
    @staticmethod
    def _deduplicate_text(text: str) -> str:
        """修复双层叠加PDF的重复字符（如 SS33 → S3, 哈哈尔尔 → 哈尔）"""
        result = []
        i = 0
        while i < len(text):
            ch = text[i]
            if i + 1 < len(text) and text[i + 1] == ch:
                result.append(ch)
                i += 2
            else:
                result.append(ch)
                i += 1
        return "".join(result)

    @staticmethod
    def _is_valid_text(text: str) -> bool:
        """判断提取的文本是否有效（非乱码）"""
        if not text or len(text.strip()) < 10:
            return False
        
        readable = 0
        for ch in text:
            if '\u4e00' <= ch <= '\u9fff':  # 中文
                readable += 1
            elif ch.isalnum():              # 英文/数字
                readable += 1
            elif ch in '，。！？、；：""''（）【】\n ':
                readable += 1
        if readable / len(text) <= 0.3:
            return False
        
        # 检测单字符行过多（竖排乱序）
        lines = [l for l in text.split('\n') if l.strip()]
        if lines:
            single_char_lines = sum(1 for l in lines if len(l.strip()) <= 1)
            if single_char_lines / len(lines) > 0.4:
                return False
        
            # 检测重复字符（双层叠加，如"哈哈尔尔"、"SS33"）
            double_count = sum(1 for l in lines if re.search(r'(.)\1', l))
            if double_count / len(lines) > 0.3:
                return False
        
        return True
    
    @staticmethod
    def parse_file(file_path):
        """根据文件扩展名选择合适的解析方法"""
        _, ext = os.path.splitext(file_path.lower())
        
        if ext == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return DocumentParser._parse_docx(file_path)
        elif ext == '.txt':
            return DocumentParser._parse_txt(file_path)
        elif ext in ['.xlsx', '.xls']:
            return DocumentParser._parse_xlsx(file_path)
        elif ext == '.csv':
            return DocumentParser._parse_csv(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff']:
            return DocumentParser._parse_image(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    @staticmethod
    def _parse_pdf(file_path):
        """解析PDF文件 - 优先文本提取，双层PDF去重修复，最后才走OCR。
        
        策略：pdfplumber 提取的文本即使不满足启发式校验也不丢弃；
        校验仅用于决定是否"额外尝试 OCR 增强"，OCR 失败或不可用时
        直接使用 pdfplumber 原始文本（带警告标记），避免用户完全无法读取文档。
        """
        content = ""
        ocr_content = ""
        deduped_content = ""

        # ═══ 方法1: pdfplumber 提取文本 ═══
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content += text + "\n"
        except Exception as e:
            print(f"[WARN] pdfplumber 解析失败: {e}")

        if not content.strip():
            # pdfplumber 完全没提取到文本 → 扫描件，走 OCR
            try:
                ocr_content = DocumentParser._ocr_pdf(file_path)
                if ocr_content.strip():
                    print(f"[INFO] OCR 提取PDF成功: {file_path}, 长度={len(ocr_content)}")
                    return ocr_content.strip()
            except Exception as e:
                print(f"[WARN] OCR 解析失败: {e}")
                import traceback
                print(traceback.format_exc())
            raise ValueError(
                "无法从PDF提取文字，文件可能是扫描图片且OCR不可用。"
                "请确保 PaddleOCR 已安装。"
            )

        # ═══ 方法1b: pdfplumber 有文本 → 校验质量 ═══
        if DocumentParser._is_valid_text(content):
            print(f"[INFO] pdfplumber 提取PDF成功: {file_path}, 长度={len(content)}")
            return content.strip()

        # 双层PDF检测：尝试去重修复
        deduped_content = DocumentParser._deduplicate_text(content)
        if DocumentParser._is_valid_text(deduped_content):
            print(f"[INFO] pdfplumber 双层PDF去重修复成功: {file_path}, 长度={len(deduped_content)}")
            return deduped_content.strip()

        # pdfplumber 文本未通过校验 → 尝试 OCR 增强（不丢弃原始文本）
        print(f"[WARN] pdfplumber 提取内容未通过校验（{len(content)}字），尝试OCR增强")
        try:
            ocr_content = DocumentParser._ocr_pdf(file_path)
        except Exception as e:
            print(f"[WARN] OCR 不可用: {e}")

        # OCR 可用且有结果 → 用 OCR
        if ocr_content.strip() and DocumentParser._is_valid_text(ocr_content):
            print(f"[INFO] OCR 增强成功，替换 pdfplumber 结果: {file_path}")
            return ocr_content.strip()

        # OCR 不可用 → 降级使用 pdfplumber 文本，带警告
        fallback = deduped_content if deduped_content.strip() else content.strip()
        print(f"[WARN] PDF 解析降级：OCR不可用，使用 pdfplumber 原始文本（{len(fallback)}字）")
        return f"[解析警告: 文本可能包含格式噪声，建议人工核对]\n\n{fallback}"
    
    @staticmethod
    def _ocr_pdf(file_path):
        """使用 OCR 识别 PDF 中的文字（PaddleOCR）"""
        from paddle_ocr import ocr_pdf
        return ocr_pdf(file_path)
    
    @staticmethod
    def _parse_docx(file_path):
        """解析Word文档（含嵌入图片OCR）"""
        content = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + "\t"
                    content += "\n"

            # 提取嵌入图片并 OCR
            import tempfile
            from paddle_ocr import ocr_image
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" not in rel.reltype:
                    continue
                try:
                    image = rel.target_part
                    suffix = os.path.splitext(image.partname)[1] or ".png"
                    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
                        f.write(image.blob)
                        tmp_path = f.name
                    ocr_text = ocr_image(tmp_path)
                    os.unlink(tmp_path)
                    if ocr_text.strip() and not ocr_text.startswith("[OCR错误"):
                        image_count += 1
                        content += "\n" + ocr_text.strip() + "\n"
                except Exception:
                    continue

            if image_count > 0 and not content.strip():
                content = f"(本文档仅有图片，已通过OCR提取 {image_count} 张图片文字)\n" + content
            return content.strip()
        except Exception as e:
            raise ValueError(f"解析Word文件失败: {str(e)}")
    
    @staticmethod
    def _parse_txt(file_path):
        """解析文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"解析文本文件失败: {str(e)}")
    
    @staticmethod
    def _parse_image(file_path):
        """解析图片文件（使用 PaddleOCR）"""
        from paddle_ocr import ocr_image
        result = ocr_image(file_path)
        if not result.strip() or result.startswith("[OCR错误"):
            raise ValueError(f"无法从图片中识别文字: {result}")
        print(f"[INFO] 图片OCR识别成功: {file_path}, 字符数={len(result)}")
        return result.strip()
    
    @staticmethod
    def _parse_xlsx(file_path):
        """解析 Excel 表格，返回结构化文本摘要（前 50 行 + 统计信息）"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("解析 Excel 需要 pandas/openpyxl，请运行: pip install pandas openpyxl")

        try:
            df = pd.read_excel(file_path, header=0)
        except Exception:
            df = pd.read_excel(file_path, header=None)

        return DocumentParser._df_to_text(df, file_path)

    @staticmethod
    def _parse_csv(file_path):
        """解析 CSV 表格，返回结构化文本摘要（前 50 行 + 统计信息）"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("解析 CSV 需要 pandas，请运行: pip install pandas")

        last_err = None
        for enc in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                df = pd.read_csv(file_path, encoding=enc)
                return DocumentParser._df_to_text(df, file_path)
            except UnicodeDecodeError as e:
                last_err = e
                continue
        raise ValueError(f"无法解析 CSV 文件编码: {file_path}（{last_err}）")

    @staticmethod
    def _df_to_text(df, file_path: str) -> str:
        """将 DataFrame 转为可读文本摘要（供 LLM 理解表格结构）"""
        lines = [f"[表格文件] {os.path.basename(file_path)}"]
        lines.append(f"行数: {len(df)}, 列数: {len(df.columns)}")
        lines.append(f"列名: {', '.join(str(c) for c in df.columns)}")
        lines.append("")

        # 数值列统计
        num_cols = df.select_dtypes(include=['number']).columns
        if len(num_cols) > 0:
            lines.append("## 数值列统计")
            for col in num_cols:
                lines.append(f"- {col}: 合计={df[col].sum():.2f}, "
                             f"均值={df[col].mean():.2f}, "
                             f"最大={df[col].max()}, "
                             f"最小={df[col].min()}")
            lines.append("")

        # 前 50 行预览
        preview_rows = min(50, len(df))
        lines.append(f"## 数据预览（前 {preview_rows} 行）")
        lines.append(df.head(preview_rows).to_string(index=True))
        return "\n".join(lines)

    @staticmethod
    def is_supported_file(file_path):
        """检查文件是否为支持的格式"""
        _, ext = os.path.splitext(file_path.lower())
        return ext in ['.pdf', '.docx', '.doc', '.txt',
                       '.xlsx', '.xls', '.csv',
                       '.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff']
